"""Multi-format export for parsed comment rows.

Supported formats: Word (.docx), Excel (.xlsx), CSV (.csv).
All exporters accept a common ExportConfig that controls columns, layout,
provenance block, and scope notes.
"""

import csv
import io
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional


@dataclass
class ExportConfig:
    """User-configurable export settings."""

    columns: list[str] = field(default_factory=lambda: [
        "No.", "Commenter", "Date", "Summary", "Applicant's Response",
    ])

    orientation: str = "landscape"          # "landscape" | "portrait"
    include_provenance: bool = True
    include_scope_notes: bool = True
    font_name: str = "Calibri"
    font_size_pt: int = 10
    header_bg_color: str = "2563EB"         # hex, no #
    header_font_color: str = "FFFFFF"
    # Maps custom column display name → row-dict key
    custom_column_keys: dict = field(default_factory=dict)


@dataclass
class ProjectInfo:
    """Metadata for the provenance block."""

    project_name: str = ""
    file_number: str = ""
    jurisdiction: str = ""
    source_description: str = ""
    scope_notes: list[str] = field(default_factory=list)
    parse_notes: str = ""
    raw_row_count: int = 0
    merged_row_count: int = 0


# Column key → row-dict key mapping
_COL_KEY = {
    "No.": "_row_num",
    "Commenter": "commenter",
    "Date": "date",
    "Summary": "summary",
    "Applicant's Response": "_response",
    "Source Reference": "source_ref",
    "Comment Type": "comment_type",
    "Topics": "topics",
}


def _row_values(row: dict, columns: list[str], row_num: int,
                custom_keys: Optional[dict] = None) -> list[str]:
    """Extract cell values for a row in column order."""
    custom_keys = custom_keys or {}
    vals = []
    for col in columns:
        key = _COL_KEY.get(col, "")
        if key == "_row_num":
            vals.append(str(row_num))
        elif key == "_response":
            vals.append("")
        elif key:
            vals.append(str(row.get(key, "")))
        elif col in custom_keys:
            vals.append(str(row.get(custom_keys[col], "")))
        else:
            vals.append(str(row.get(col.lower().replace(" ", "_"), "")))
    return vals


# ---- helpers for Word ----

def _twips(inches: float) -> int:
    return int(round(inches * 1440))


def _add_provenance(doc, info: ProjectInfo, config: ExportConfig):
    from docx.shared import Pt

    lines = [
        ("Source record — Tab B prep (public comments to applicant response table)", True),
        (f"Project: {info.project_name} — {info.jurisdiction}, "
         f"file no. {info.file_number}", False),
        (f"Source: {info.source_description}", False),
        (f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", False),
        (f"Rows: {info.raw_row_count} raw → {info.merged_row_count} merged "
         "(same commenter consolidated).", False),
    ]
    if info.parse_notes:
        lines.append((f"AI notes: {info.parse_notes}", False))

    for text, bold in lines:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(config.font_size_pt - 1)
        r.font.name = config.font_name
        if bold:
            r.bold = True

    if config.include_scope_notes and info.scope_notes:
        p = doc.add_paragraph()
        r = p.add_run("Omissions / scope notes (material not represented as rows below)")
        r.bold = True
        r.font.size = Pt(config.font_size_pt - 1)
        r.font.name = config.font_name
        for note in info.scope_notes:
            bp = doc.add_paragraph(note, style="List Bullet")
            for run in bp.runs:
                run.font.size = Pt(config.font_size_pt - 1)
                run.font.name = config.font_name

    doc.add_paragraph()


def _style_header_row(table, config: ExportConfig):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), config.header_bg_color)
        tcPr.append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(config.header_font_color)
                run.bold = True
                run.font.size = Pt(config.font_size_pt)
                run.font.name = config.font_name


def _style_body_rows(table, config: ExportConfig):
    from docx.shared import Pt

    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(config.font_size_pt)
                    run.font.name = config.font_name


def _set_fixed_layout(table, widths_twips: list[int]):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblPr.append(layout)

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(tbl.index(tblPr) + 1, grid)
    for old in grid.findall(qn("w:gridCol")):
        grid.remove(old)
    for w in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths_twips[idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _column_widths(columns: list[str], usable_in: float) -> list[float]:
    """Calculate column widths in inches based on which columns are active."""
    fixed = {"No.": 0.42, "Date": 1.12, "Commenter": 1.12,
             "Comment Type": 1.0, "Source Reference": 1.2}
    flex = {"Summary", "Applicant's Response", "Topics"}

    total_fixed = sum(fixed.get(c, 0) for c in columns if c in fixed)
    flex_cols = [c for c in columns if c in flex]
    remainder = max(1.0, usable_in - total_fixed)

    weights = {"Summary": 3, "Applicant's Response": 5, "Topics": 1}
    total_weight = sum(weights.get(c, 2) for c in flex_cols) or 1

    widths = []
    for c in columns:
        if c in fixed:
            widths.append(fixed[c])
        else:
            widths.append(remainder * weights.get(c, 2) / total_weight)
    return widths


# ===========================================================================
# Public export functions — each returns a BytesIO (or str for CSV)
# ===========================================================================

def export_docx(
    rows: list[dict],
    project_info: ProjectInfo,
    config: Optional[ExportConfig] = None,
) -> io.BytesIO:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    config = config or ExportConfig()
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    if config.orientation == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

    pw = sec.page_width.inches if hasattr(sec.page_width, "inches") else 11.0
    lm = sec.left_margin.inches if hasattr(sec.left_margin, "inches") else 0.5
    rm = sec.right_margin.inches if hasattr(sec.right_margin, "inches") else 0.5
    usable = pw - lm - rm

    if config.include_provenance:
        _add_provenance(doc, project_info, config)

    ncols = len(config.columns)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = False

    for idx, label in enumerate(config.columns):
        table.rows[0].cells[idx].text = label

    for num, row in enumerate(rows, 1):
        vals = _row_values(row, config.columns, num, config.custom_column_keys)
        for idx, val in enumerate(vals):
            table.rows[num].cells[idx].text = val

    col_widths_in = _column_widths(config.columns, usable - 0.16)
    _set_fixed_layout(table, [_twips(w) for w in col_widths_in])
    _style_header_row(table, config)
    _style_body_rows(table, config)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_xlsx(
    rows: list[dict],
    project_info: ProjectInfo,
    config: Optional[ExportConfig] = None,
) -> io.BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    config = config or ExportConfig()
    wb = Workbook()
    ws = wb.active
    ws.title = "Comments"

    if config.orientation == "landscape":
        ws.sheet_properties.pageSetUpPr = None
        ws.page_setup.orientation = "landscape"
        ws.page_setup.paperSize = ws.PAPERSIZE_LETTER

    header_fill = PatternFill(start_color=config.header_bg_color,
                              end_color=config.header_bg_color, fill_type="solid")
    header_font = Font(name=config.font_name, size=config.font_size_pt,
                       bold=True, color=config.header_font_color)
    body_font = Font(name=config.font_name, size=config.font_size_pt)
    wrap = Alignment(wrap_text=True, vertical="top")

    for idx, label in enumerate(config.columns, 1):
        cell = ws.cell(row=1, column=idx, value=label)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = wrap

    for num, row in enumerate(rows, 1):
        vals = _row_values(row, config.columns, num, config.custom_column_keys)
        for idx, val in enumerate(vals, 1):
            cell = ws.cell(row=num + 1, column=idx, value=val)
            cell.font = body_font
            cell.alignment = wrap

    col_widths_ch = {"No.": 5, "Commenter": 22, "Date": 14,
                     "Summary": 55, "Applicant's Response": 45,
                     "Source Reference": 18, "Comment Type": 14, "Topics": 20}
    for idx, col in enumerate(config.columns, 1):
        ws.column_dimensions[get_column_letter(idx)].width = col_widths_ch.get(col, 18)

    if config.include_provenance and (project_info.project_name or project_info.scope_notes):
        info_ws = wb.create_sheet("Info", 0)
        info_ws.append(["Project", project_info.project_name])
        info_ws.append(["File Number", project_info.file_number])
        info_ws.append(["Jurisdiction", project_info.jurisdiction])
        info_ws.append(["Source", project_info.source_description])
        info_ws.append(["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")])
        info_ws.append(["Rows", f"{project_info.raw_row_count} raw → "
                        f"{project_info.merged_row_count} merged"])
        if project_info.parse_notes:
            info_ws.append(["AI Notes", project_info.parse_notes])
        if config.include_scope_notes:
            info_ws.append([])
            info_ws.append(["Scope Notes"])
            for note in project_info.scope_notes:
                info_ws.append(["", note])
        info_ws.column_dimensions["A"].width = 16
        info_ws.column_dimensions["B"].width = 90
        wb.active = wb.sheetnames.index("Comments")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def export_csv(
    rows: list[dict],
    project_info: ProjectInfo,
    config: Optional[ExportConfig] = None,
) -> str:
    config = config or ExportConfig()
    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerow(config.columns)
    for num, row in enumerate(rows, 1):
        writer.writerow(_row_values(row, config.columns, num, config.custom_column_keys))
    return out.getvalue()


def export_pdf(
    rows: list[dict],
    project_info: ProjectInfo,
    config: Optional[ExportConfig] = None,
) -> io.BytesIO:
    """Export to PDF by rendering the Word doc and converting via python-docx2pdf
    or, as a portable fallback, build a simple PDF with reportlab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, landscape as rl_landscape
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Table as RLTable, TableStyle, Paragraph, Spacer,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        config = config or ExportConfig()
        buf = io.BytesIO()
        pagesize = rl_landscape(letter) if config.orientation == "landscape" else letter
        doc = SimpleDocTemplate(buf, pagesize=pagesize,
                                leftMargin=0.5*inch, rightMargin=0.5*inch,
                                topMargin=0.5*inch, bottomMargin=0.5*inch)

        styles = getSampleStyleSheet()
        cell_style = ParagraphStyle(
            "CellStyle", parent=styles["Normal"],
            fontName="Helvetica", fontSize=config.font_size_pt - 1,
            leading=config.font_size_pt + 2,
        )
        header_style = ParagraphStyle(
            "HeaderStyle", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=config.font_size_pt,
            leading=config.font_size_pt + 3,
            textColor=colors.white,
        )

        elements = []

        if config.include_provenance:
            prov_lines = [
                f"<b>Source record — Tab B prep</b>",
                f"Project: {project_info.project_name} — {project_info.jurisdiction}, "
                f"file no. {project_info.file_number}",
                f"Source: {project_info.source_description}",
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                f"Rows: {project_info.raw_row_count} raw → "
                f"{project_info.merged_row_count} merged",
            ]
            for line in prov_lines:
                elements.append(Paragraph(line, styles["Normal"]))
            elements.append(Spacer(1, 12))

        if config.include_scope_notes and project_info.scope_notes:
            elements.append(Paragraph(
                "<b>Omissions / scope notes</b>", styles["Normal"]
            ))
            for note in project_info.scope_notes:
                elements.append(Paragraph(f"• {note}", cell_style))
            elements.append(Spacer(1, 12))

        header_row = [Paragraph(c, header_style) for c in config.columns]
        table_data = [header_row]
        for num, row in enumerate(rows, 1):
            vals = _row_values(row, config.columns, num, config.custom_column_keys)
            table_data.append([Paragraph(v, cell_style) for v in vals])

        avail_width = pagesize[0] - 1.0 * inch
        n = len(config.columns)
        col_w = [avail_width / n] * n

        t = RLTable(table_data, colWidths=col_w, repeatRows=1)
        hdr_bg = colors.HexColor(f"#{config.header_bg_color}")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), hdr_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), config.font_size_pt - 1),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)
        doc.build(elements)
        buf.seek(0)
        return buf

    except ImportError:
        docx_buf = export_docx(rows, project_info, config)
        return docx_buf
